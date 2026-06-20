# build_report_docx.py
# ============================================================
# Word 报告生成脚本：读取模板、插入文字/表格/图片，并保存项目报告
# ============================================================
#
# 这个脚本不参与模型训练，它的作用是把实验结果自动写成 .docx 报告。
# 可以把它理解成“自动排版员”：它读取模板和 report_assets 中的图表/指标，
# 然后按章节生成封面、摘要、数据集介绍、模型说明、实验结果和参考资料。

# 从 pathlib 导入 Path；Path 比普通字符串路径更适合拼接和读取文件。
from pathlib import Path

# 导入 json，用来读取 report_assets/metrics.json 中的实验指标。
import json

# 从 python-docx 导入 Document，用来打开、编辑和保存 Word 文档。
from docx import Document

# 导入分节相关枚举；当前脚本没有主动分节，但保留导入不影响运行。
from docx.enum.section import WD_SECTION_START

# 导入表格垂直对齐和表格整体对齐枚举。
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

# 导入段落对齐和换行枚举；WD_BREAK 当前未使用，但保留不改变原依赖结构。
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

# 导入 OxmlElement，用于设置 python-docx 高层 API 没覆盖的表格边框和底纹。
from docx.oxml import OxmlElement

# 导入 qn，用来把 w:xxx 这类 Word XML 名称转换成完整命名空间名称。
from docx.oxml.ns import qn

# 导入尺寸和颜色工具；Inches 表示英寸，Pt 表示磅，RGBColor 表示 RGB 颜色。
from docx.shared import Inches, Pt, RGBColor


# BASE_DIR 是当前脚本所在目录，也就是项目根目录。
BASE_DIR = Path(__file__).resolve().parent

# TEMPLATE 是 Word 模板路径；报告会基于这个模板创建。
TEMPLATE = BASE_DIR / "报告模板(1).docx"

# OUTPUT 是最终生成的报告文件路径。
OUTPUT = BASE_DIR / "进食声音识别项目报告.docx"

# ASSET_DIR 是报告图片和指标文件所在目录。
ASSET_DIR = BASE_DIR / "report_assets"

# METRICS 读取 metrics.json，里面保存类别、准确率、混淆矩阵分析等结果。
METRICS = json.loads((ASSET_DIR / "metrics.json").read_text(encoding="utf-8"))

# CHECKPOINT 保存当前 best_model.pth 对应的最优轮次和验证准确率。
CHECKPOINT = METRICS["checkpoint"]

# EVALUATION 保存本次评估的验证集规模、分类报告和混淆分析结果。
EVALUATION = METRICS["evaluation"]

# REPORT 是 sklearn classification_report 的结构化结果。
REPORT = EVALUATION["classification_report"]

# TRAINING_RUN 保存最新逐轮训练历史的摘要。
TRAINING_RUN = METRICS.get("training_log_runs", [{}])[0]


# 定义百分比格式化函数；支持 0.9445 这类小数和 94.45 这类百分数。
def fmt_pct(value):
    # 转成浮点数，便于统一处理。
    value = float(value)

    # 小于等于 1 的值按小数准确率处理。
    if value <= 1:
        value *= 100

    # 保留两位小数，和报告表格中的格式一致。
    return f"{value:.2f}%"


# 定义类别召回率格式化函数。
def class_recall(name):
    # classification_report 中每个类别的 recall 就是该类验证样本的召回率。
    return fmt_pct(REPORT[name]["recall"])


# 定义类别精确率格式化函数。
def class_precision(name):
    # classification_report 中每个类别的 precision 反映被预测为该类时的可靠程度。
    return fmt_pct(REPORT[name]["precision"])


# 定义类别 F1 格式化函数。
def class_f1(name):
    # F1 综合精确率和召回率。
    return fmt_pct(REPORT[name]["f1-score"])


# 定义从 top_confusions 中读取指定误分数量的函数。
def confusion_count(true_name, pred_name):
    # 遍历最新误分列表，查找指定 true -> pred 组合。
    for item in EVALUATION["top_confusions"]:
        if item["true"] == true_name and item["pred"] == pred_name:
            return item["count"]

    # 未进入 top_confusions 的组合按 0 处理，避免报告误写旧数字。
    return 0


# CN_FONT 是正文中文字体。
CN_FONT = "宋体"

# HEAD_FONT 是标题中文字体。
HEAD_FONT = "黑体"

# EN_FONT 是英文字体。
EN_FONT = "Times New Roman"

# BLUE 是一级标题常用蓝色。
BLUE = RGBColor(46, 116, 181)

# DARK 是二级标题常用深蓝色。
DARK = RGBColor(31, 77, 120)


# CLASS_CN 把英文类别名映射成中文说明，方便报告表格展示。
CLASS_CN = {
    # aloe 对应芦荟。
    "aloe": "芦荟",

    # burger 对应汉堡。
    "burger": "汉堡",

    # cabbage 对应卷心菜。
    "cabbage": "卷心菜",

    # candied_fruits 对应蜜饯或糖葫芦类。
    "candied_fruits": "蜜饯/糖葫芦类",

    # carrots 对应胡萝卜。
    "carrots": "胡萝卜",

    # chips 对应薯片。
    "chips": "薯片",

    # chocolate 对应巧克力。
    "chocolate": "巧克力",

    # drinks 对应饮品。
    "drinks": "饮品",

    # fries 对应薯条。
    "fries": "薯条",

    # grapes 对应葡萄。
    "grapes": "葡萄",

    # gummies 对应软糖。
    "gummies": "软糖",

    # ice-cream 对应冰淇淋。
    "ice-cream": "冰淇淋",

    # jelly 对应果冻。
    "jelly": "果冻",

    # noodles 对应面条。
    "noodles": "面条",

    # pickles 对应腌菜/泡菜。
    "pickles": "腌菜/泡菜",

    # pizza 对应披萨。
    "pizza": "披萨",

    # ribs 对应排骨。
    "ribs": "排骨",

    # salmon 对应三文鱼。
    "salmon": "三文鱼",

    # soup 对应汤类。
    "soup": "汤类",

    # wings 对应鸡翅。
    "wings": "鸡翅",
}


# 定义清空模板正文的函数；保留节属性，删除已有段落和表格。
def clear_body(doc):
    # 取得文档 body 对应的底层 XML 元素。
    body = doc._body._element

    # list(body) 复制一份子节点列表，避免遍历时删除导致迭代混乱。
    for child in list(body):
        # Word 的 sectPr 节属性不能随便删，否则页边距等设置可能丢失。
        if child.tag != qn("w:sectPr"):
            # 删除模板原有正文元素。
            body.remove(child)


# 定义设置 run 字体的函数；run 是 Word 中一小段连续文字。
def set_run_font(run, font=CN_FONT, size=None, bold=None, color=None):
    # 设置西文字体名称；python-docx 的 name 主要影响西文。
    run.font.name = EN_FONT

    # 设置中文字体 eastAsia，保证中文不是默认字体。
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font)

    # 设置 ASCII 字体。
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), EN_FONT)

    # 设置高 ANSI 字体。
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), EN_FONT)

    # 如果传入 size，就设置字号。
    if size is not None:
        # Pt(size) 表示以磅为单位的字号。
        run.font.size = Pt(size)

    # 如果传入 bold，就设置是否加粗。
    if bold is not None:
        # True 表示加粗，False 表示取消加粗。
        run.bold = bold

    # 如果传入 color，就设置字体颜色。
    if color is not None:
        # run.font.color.rgb 接收 RGBColor 对象。
        run.font.color.rgb = color


# 定义段落格式设置函数；统一首行缩进、行距和段前段后距离。
def set_paragraph_format(paragraph, first_line=True, line_spacing=1.5, before=0, after=6):
    # paragraph_format 是当前段落的格式对象。
    fmt = paragraph.paragraph_format

    # 根据 first_line 决定是否设置首行缩进；正文通常首行缩进 24 磅。
    fmt.first_line_indent = Pt(24) if first_line else Pt(0)

    # 设置行距，例如 1.5 倍行距。
    fmt.line_spacing = line_spacing

    # 设置段前距离。
    fmt.space_before = Pt(before)

    # 设置段后距离。
    fmt.space_after = Pt(after)


# 定义添加居中文本段落的函数，常用于封面标题。
def add_centered(doc, text, size=14, bold=False, font=CN_FONT, after=6):
    # 新增一个段落。
    p = doc.add_paragraph()

    # 设置段落水平居中。
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置段落格式；居中标题不需要首行缩进。
    set_paragraph_format(p, first_line=False, after=after)

    # 在段落里添加文字 run。
    r = p.add_run(text)

    # 设置文字字体、字号和是否加粗。
    set_run_font(r, font=font, size=size, bold=bold)

    # 返回段落对象，方便调用者继续处理。
    return p


# 定义添加正文段落的函数；支持用 \n 在同一段内换行。
def add_body(doc, text, after=6):
    # 新增一个正文段落。
    p = doc.add_paragraph()

    # 设置正文段落格式，默认首行缩进。
    set_paragraph_format(p, first_line=True, after=after)

    # 按换行符拆分文本，逐段加入同一个段落。
    for part in text.split("\n"):
        # 如果当前片段不是空字符串，就添加文字。
        if part:
            # 添加当前片段到段落。
            r = p.add_run(part)

            # 设置正文文字字体和字号。
            set_run_font(r, CN_FONT, 12)

        # 如果当前片段不是最后一个片段，就添加换行。
        if part != text.split("\n")[-1]:
            # add_break 会在 Word 段落内部插入换行符。
            p.add_run().add_break()

    # 返回正文段落对象。
    return p


# 定义添加标题段落的函数；level=1 表示一级标题，level=2 表示二级标题。
def add_heading(doc, text, level=1):
    # 新增标题段落。
    p = doc.add_paragraph()

    # 设置标题段落格式；标题不需要首行缩进。
    set_paragraph_format(p, first_line=False, before=10 if level == 1 else 6, after=6)

    # 如果是一级标题，就使用更大字号和蓝色。
    if level == 1:
        # 一级标题字号 14。
        size = 14

        # 一级标题颜色使用 BLUE。
        color = BLUE

    # 否则按二级标题处理。
    else:
        # 二级标题字号 12。
        size = 12

        # 二级标题颜色使用 DARK。
        color = DARK

    # 添加标题文字。
    r = p.add_run(text)

    # 设置标题字体、字号、加粗和颜色。
    set_run_font(r, HEAD_FONT, size, True, color)

    # 返回标题段落对象。
    return p


# 定义添加图表标题/说明文字的函数。
def add_caption(doc, text):
    # 新增说明段落。
    p = doc.add_paragraph()

    # 设置说明文字居中。
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置说明段落格式，不需要首行缩进。
    set_paragraph_format(p, first_line=False, after=8)

    # 添加说明文字。
    r = p.add_run(text)

    # 设置说明文字字体和字号。
    set_run_font(r, CN_FONT, 12)

    # 返回说明段落对象。
    return p


# 定义设置单元格文字的函数；用于统一表格风格。
def set_cell_text(cell, text, bold=False, font_size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    # 先清空单元格原有文字。
    cell.text = ""

    # 获取单元格中的第一个段落。
    p = cell.paragraphs[0]

    # 设置段落水平对齐方式。
    p.alignment = align

    # 表格中段后距离设为 0，避免行高过大。
    p.paragraph_format.space_after = Pt(0)

    # 表格中文字行距设为 1.15。
    p.paragraph_format.line_spacing = 1.15

    # 添加单元格文字，并转成字符串避免数字类型报错。
    r = p.add_run(str(text))

    # 设置单元格文字字体、字号和加粗。
    set_run_font(r, CN_FONT, font_size, bold)

    # 设置单元格垂直居中。
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# 定义设置单元格底纹颜色的函数。
def set_cell_shading(cell, fill):
    # 获取单元格属性 XML。
    tc_pr = cell._tc.get_or_add_tcPr()

    # 查找已有底纹元素。
    shd = tc_pr.find(qn("w:shd"))

    # 如果没有底纹元素，就创建一个。
    if shd is None:
        # 创建 Word XML 的 w:shd 元素。
        shd = OxmlElement("w:shd")

        # 把底纹元素加入单元格属性。
        tc_pr.append(shd)

    # 设置底纹填充色，例如 E8EEF5。
    shd.set(qn("w:fill"), fill)


# 定义设置单元格边框的函数；kwargs 指定各边框样式。
def set_cell_border(cell, **kwargs):
    # 获取单元格底层 XML。
    tc = cell._tc

    # 获取或创建单元格属性。
    tc_pr = tc.get_or_add_tcPr()

    # 查找已有边框集合。
    borders = tc_pr.first_child_found_in("w:tcBorders")

    # 如果没有边框集合，就创建一个。
    if borders is None:
        # 创建边框集合元素。
        borders = OxmlElement("w:tcBorders")

        # 添加到单元格属性中。
        tc_pr.append(borders)

    # 遍历 Word 支持的边框位置。
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        # 只有调用者传入的边框才设置。
        if edge in kwargs:
            # 拼出 Word XML 标签名，例如 w:top。
            tag = "w:{}".format(edge)

            # 查找这个边框元素是否已经存在。
            element = borders.find(qn(tag))

            # 如果不存在，就新建。
            if element is None:
                # 创建对应边框元素。
                element = OxmlElement(tag)

                # 添加到边框集合。
                borders.append(element)

            # 遍历边框的属性字典，例如 val、sz、color。
            for key, value in kwargs[edge].items():
                # 设置边框属性值；Word XML 属性也需要带命名空间。
                element.set(qn("w:{}".format(key)), str(value))


# 定义统一表格样式的函数。
def style_table(table, header=True):
    # 设置表格整体居中。
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 启用自动适应列宽。
    table.autofit = True

    # 遍历表格中的每一行，同时得到行号。
    for row_idx, row in enumerate(table.rows):
        # 遍历当前行中的每个单元格。
        for cell in row.cells:
            # 如果 header=True 且当前是第一行，就设置表头底纹。
            if header and row_idx == 0:
                # 给表头单元格填充浅蓝灰色。
                set_cell_shading(cell, "E8EEF5")

            # 给单元格设置四周边框。
            set_cell_border(
                # 当前单元格对象。
                cell,

                # 上边框样式。
                top={"val": "single", "sz": "6", "color": "7F7F7F"},

                # 下边框样式。
                bottom={"val": "single", "sz": "6", "color": "7F7F7F"},

                # 左边框样式。
                left={"val": "single", "sz": "4", "color": "BFBFBF"},

                # 右边框样式。
                right={"val": "single", "sz": "4", "color": "BFBFBF"},
            )


# 定义添加表格的通用函数。
def add_table(doc, headers, rows, widths=None, font_size=10):
    # 新建一个表格，先创建 1 行表头，列数等于 headers 长度。
    table = doc.add_table(rows=1, cols=len(headers))

    # 先设置一次表格样式。
    style_table(table, header=True)

    # 遍历表头文字和列号。
    for idx, header in enumerate(headers):
        # 设置表头单元格文字，表头加粗。
        set_cell_text(table.rows[0].cells[idx], header, bold=True, font_size=font_size)

        # 如果传入列宽，就设置当前表头单元格宽度。
        if widths:
            # Inches 将数字转换为 Word 需要的英寸长度。
            table.rows[0].cells[idx].width = Inches(widths[idx])

    # 遍历数据行。
    for row in rows:
        # 新增一行，并取得这一行的单元格列表。
        cells = table.add_row().cells

        # 遍历当前行的每个值。
        for idx, value in enumerate(row):
            # 如果是最后一列且列数较多，通常是说明文字，使用左对齐更易读。
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == len(row) - 1 and len(row) > 3 else WD_ALIGN_PARAGRAPH.CENTER

            # 设置当前单元格文字。
            set_cell_text(cells[idx], value, font_size=font_size, align=align)

            # 如果传入列宽，就设置当前单元格宽度。
            if widths:
                # 用对应的 widths[idx] 设置列宽。
                cells[idx].width = Inches(widths[idx])

    # 数据填充后再统一应用一次样式，确保新增行也有边框。
    style_table(table, header=True)

    # 表格后添加一个空段落，让后续内容与表格之间有间隔。
    doc.add_paragraph()

    # 返回表格对象。
    return table


# 定义添加图片和图注的函数。
def add_picture(doc, image_name, width, caption):
    # 新增一个段落专门放图片。
    p = doc.add_paragraph()

    # 图片段落居中。
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置图片段落格式，不首行缩进，段后距离较小。
    set_paragraph_format(p, first_line=False, after=2)

    # 添加一个 run，用来承载图片。
    run = p.add_run()

    # 插入图片，路径来自 ASSET_DIR，宽度用英寸控制。
    run.add_picture(str(ASSET_DIR / image_name), width=Inches(width))

    # 图片下面添加图注。
    add_caption(doc, caption)


# 定义添加项目符号列表的函数。
def add_bullets(doc, items):
    # 遍历每个列表项文字。
    for item in items:
        # 尝试使用 Word 模板中的 List Bullet 样式。
        try:
            # 创建带项目符号样式的段落。
            p = doc.add_paragraph(style="List Bullet")

        # 如果模板没有这个样式，就退回普通段落。
        except KeyError:
            # 创建普通段落。
            p = doc.add_paragraph()

        # 设置列表段落格式，不首行缩进。
        set_paragraph_format(p, first_line=False, after=3)

        # 添加列表项文字。
        r = p.add_run(item)

        # 设置列表项字体。
        set_run_font(r, CN_FONT, 12)


# 定义文档全局样式设置函数。
def setup_styles(doc):
    # 取得第一个节；通常整篇报告只用一个节。
    section = doc.sections[0]

    # 设置上页边距为 1 英寸。
    section.top_margin = Inches(1)

    # 设置下页边距为 1 英寸。
    section.bottom_margin = Inches(1)

    # 设置左页边距为 1 英寸。
    section.left_margin = Inches(1)

    # 设置右页边距为 1 英寸。
    section.right_margin = Inches(1)

    # 获取 Normal 正文样式。
    normal = doc.styles["Normal"]

    # 设置正文西文字体。
    normal.font.name = EN_FONT

    # 设置正文中文字体。
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    # 设置正文 ASCII 字体。
    normal._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)

    # 设置正文高 ANSI 字体。
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)

    # 设置正文默认字号。
    normal.font.size = Pt(12)


# 定义封面生成函数。
def add_cover(doc):
    # 添加两个空段落，让封面标题往下移动。
    for _ in range(2):
        # 新增空段落。
        doc.add_paragraph()

    # 添加课程名称。
    add_centered(doc, "语音识别技术与应用课程", size=20, bold=True, font=HEAD_FONT, after=10)

    # 添加报告类型标题。
    add_centered(doc, "期末项目报告", size=24, bold=True, font=HEAD_FONT, after=22)

    # 添加项目名称。
    add_centered(doc, "项目名称：基于 CNN 的进食声音分类识别系统", size=15, bold=True, font=CN_FONT, after=28)

    # fields 保存封面需要填写的信息行。
    fields = [
        # 学院字段。
        "学院：昆明城市学院",

        # 专业字段。
        "专业：_____________________________",

        # 班级字段。
        "班级：_____________________________",

        # 组长字段。
        "组长姓名：_________________________",

        # 小组成员字段。
        "小组成员：_________________________",
    ]

    # 逐行添加封面字段。
    for field in fields:
        # 新增段落。
        p = doc.add_paragraph()

        # 字段段落左对齐。
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 设置左缩进，使字段出现在封面中间偏左。
        p.paragraph_format.left_indent = Inches(1.4)

        # 设置字段行之间的间距。
        p.paragraph_format.space_after = Pt(10)

        # 添加字段文字。
        r = p.add_run(field)

        # 设置字段字体和字号。
        set_run_font(r, CN_FONT, 14)

    # 再添加两个空段落，让日期下移。
    for _ in range(2):
        # 新增空段落。
        doc.add_paragraph()

    # 添加日期。
    add_centered(doc, "2026年6月20日", size=12, font=CN_FONT)

    # 封面结束后插入分页符。
    doc.add_page_break()


# 定义摘要章节生成函数。
def add_abstract(doc):
    # 添加“摘要”一级标题。
    add_heading(doc, "摘要", 1)

    # 添加摘要第一段，概括任务、数据处理和模型输入。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 摘要正文第一段。
        "本项目面向进食声音识别任务，使用 Kaggle Eating Sound Collection 数据集构建 20 类食物/进食声音分类系统。项目以音频文件为输入，统一完成单声道转换、采样率重采样、3 秒截断或补零、Mel 频谱图提取和归一化处理，再将 Mel 频谱图作为单通道图像输入卷积神经网络进行分类。",
    )

    # 添加摘要第二段，说明最终模型和效果。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 摘要正文第二段。
        f"当前最终模型采用 {CHECKPOINT['arch']} 结构，在验证集 {EVALUATION['val_samples']} 条样本上取得 {fmt_pct(EVALUATION['accuracy'])} 的准确率，最优 checkpoint 保存于第 {CHECKPOINT['epoch']} 轮。与早期 cnn_v1 留痕实验相比，最终模型通过更深的双卷积模块、SiLU 激活、Dropout2d 正则化、类别权重和 SpecAugment 数据增强显著提升了识别效果。项目代码包含训练、预测、数据缓存和模型保存功能，能够支持单个音频文件的 Top-5 类别预测。",
    )

    # 新增关键词段落。
    p = doc.add_paragraph()

    # 设置关键词段落格式。
    set_paragraph_format(p, first_line=True, after=8)

    # 添加“关键词：”四个字。
    r = p.add_run("关键词：")

    # 设置“关键词：”为黑体加粗。
    set_run_font(r, HEAD_FONT, 12, True)

    # 添加具体关键词内容。
    r = p.add_run("进食声音识别；Mel 频谱图；卷积神经网络；PyTorch；音频分类")

    # 设置关键词内容为正文样式。
    set_run_font(r, CN_FONT, 12)


# 定义数据集章节生成函数。
def add_dataset_section(doc):
    # 添加数据集介绍一级标题。
    add_heading(doc, "1 数据集介绍", 1)

    # 添加数据来源二级标题。
    add_heading(doc, "1.1 数据来源与任务定义", 2)

    # 添加数据来源说明段。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 数据来源正文。
        "本项目数据集来源为 Kaggle 平台的 Eating Sound Collection 数据集，下载地址为：https://www.kaggle.com/datasets/mashijie/eating-sound-collection?resource=download。项目本地数据存放在 archive/clips_rd 目录下，按照类别文件夹组织音频文件。任务目标是从短时进食声音中识别对应的食物或进食类别，属于多类别音频分类问题。",
    )

    # 添加数据规模说明段。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 数据规模正文。
        f"数据集共包含 {len(METRICS['classes'])} 个类别、{EVALUATION['total_samples']} 条音频样本。类别包括 aloe、burger、cabbage、candied_fruits、carrots、chips、chocolate、drinks、fries、grapes、gummies、ice-cream、jelly、noodles、pickles、pizza、ribs、salmon、soup 和 wings。项目使用分层划分方式按照约 8:2 的比例划分训练集和验证集，当前代码对应训练样本 {EVALUATION['train_samples']} 条、验证样本 {EVALUATION['val_samples']} 条。",
    )

    # 创建类别表格的数据行列表。
    rows = []

    # 遍历 metrics 中的类别列表，并从 1 开始编号。
    for idx, name in enumerate(METRICS["classes"], 1):
        # 添加一行：序号、英文类别、中文说明、样本数。
        rows.append([idx, name, CLASS_CN.get(name, ""), METRICS["class_counts"][name]])

    # 添加表 1 标题。
    add_caption(doc, "表1 数据集类别与样本数量")

    # 添加类别统计表格。
    add_table(doc, ["序号", "英文类别", "中文说明", "样本数"], rows, widths=[0.6, 1.7, 2.4, 1.0], font_size=9.5)

    # 添加类别分布图。
    add_picture(doc, "class_distribution.png", 6.2, "图1 数据集类别分布")

    # 添加预处理二级标题。
    add_heading(doc, "1.2 音频预处理与特征提取", 2)

    # 添加预处理流程说明段。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 预处理说明正文。
        "项目在 dataset.py 中实现音频预处理流程。首先将多通道音频平均为单声道；若原始采样率不等于目标采样率，则重采样到 22050 Hz；随后将每条音频统一裁剪或补零到 3 秒。完成长度统一后，使用 torchaudio.transforms.MelSpectrogram 将波形转换为 Mel 频谱图，再通过 AmplitudeToDB 转换到 dB 尺度。训练和验证阶段均将 Mel 特征按样本自身的最小值与最大值归一化到 [0,1] 区间。",
    )

    # 添加缓存和增强说明段。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 缓存与增强正文。
        "为提升训练效率，PrecomputedMelDataset 会在首次运行时把每条音频的 dB Mel 频谱图缓存为 .pt 文件，后续训练直接从 mel_cache 读取缓存特征。训练集额外使用 TimeMasking、FrequencyMasking 和时间轴随机平移增强，验证集仅做归一化，不引入随机增强。",
    )

    # 定义预处理参数表数据。
    preprocess_rows = [
        # 采样率行。
        ["采样率", "22050 Hz"],

        # 音频长度行。
        ["音频长度", "3.0 秒，不足补零，超出截断"],

        # Mel 滤波器数量行。
        ["Mel 滤波器数量", "128"],

        # FFT 点数行。
        ["FFT 点数", "1024"],

        # 帧移行。
        ["帧移 hop_length", "512"],

        # 频率范围行。
        ["频率范围", "20 Hz - 8000 Hz"],

        # 幅度转换行。
        ["幅度转换", "AmplitudeToDB(top_db=80)"],

        # 数据增强行。
        ["数据增强", "TimeMasking=20、FrequencyMasking=10、时间平移最大 8 帧"],
    ]

    # 添加表 2 标题。
    add_caption(doc, "表2 音频预处理与特征提取参数")

    # 添加预处理参数表。
    add_table(doc, ["参数", "取值/说明"], preprocess_rows, widths=[1.8, 4.5], font_size=10)


# 定义模型章节生成函数。
def add_model_section(doc):
    # 添加模型搭建一级标题。
    add_heading(doc, "2 模型搭建", 1)

    # 添加总体思路二级标题。
    add_heading(doc, "2.1 模型总体思路", 2)

    # 添加模型总体思路正文。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 模型思路正文。
        "项目将 Mel 频谱图视为一张单通道图像，使用二维卷积网络提取频率轴和时间轴上的局部模式。相比直接处理原始波形，Mel 频谱图能更直观地表达咀嚼、吞咽、脆裂、液体流动等进食声音的能量分布变化，适合使用 CNN 进行模式识别。",
    )

    # 添加输入输出说明正文。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 输入输出说明。
        "当前配置 MODEL_ARCH 为 cnn_v2。模型输入形状为 (B,1,128,T)，其中 B 为批量大小，128 为 Mel 频带数量，T 为由 3 秒音频和 hop_length 决定的时间帧数。模型输出为 20 维 logits，对应 20 个食物声音类别。",
    )

    # 添加模型结构图。
    add_picture(doc, "model_architecture.png", 6.3, "图2 cnn_v2 模型结构示意图")

    # 添加 cnn_v2 结构说明二级标题。
    add_heading(doc, "2.2 cnn_v2 结构说明", 2)

    # 添加 cnn_v2 文字说明。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # cnn_v2 结构说明。
        "cnn_v2 由 4 个 DoubleConvBlock 组成，每个阶段包含两层 3×3 卷积、BatchNorm2d 和 SiLU 激活函数，并在阶段末尾使用 MaxPool2d 下采样。后两层卷积阶段加入更强的 Dropout2d 正则化，以降低过拟合风险。特征提取后，模型使用 AdaptiveAvgPool2d 将二维特征图压缩为 1×1，再通过全连接分类器输出最终类别。",
    )

    # 定义模型结构表格数据。
    model_rows = [
        # 输入层说明。
        ["输入层", "单通道 Mel 频谱图", "(B,1,128,T)"],

        # 第一阶段说明。
        ["Stage 1", "DoubleConvBlock，1→32，Dropout2d=0.05", "下采样"],

        # 第二阶段说明。
        ["Stage 2", "DoubleConvBlock，32→64，Dropout2d=0.05", "下采样"],

        # 第三阶段说明。
        ["Stage 3", "DoubleConvBlock，64→128，Dropout2d=0.10", "下采样"],

        # 第四阶段说明。
        ["Stage 4", "DoubleConvBlock，128→256，Dropout2d=0.10", "下采样"],

        # 全局池化说明。
        ["全局池化", "AdaptiveAvgPool2d((1,1))", "固定输出维度"],

        # 分类器说明。
        ["分类器", "Dropout(0.35) + Linear(256,256) + SiLU + Dropout(0.175) + Linear(256,20)", "20 类输出"],
    ]

    # 添加表 3 标题。
    add_caption(doc, "表3 cnn_v2 模型结构")

    # 添加模型结构表。
    add_table(doc, ["模块", "结构", "作用"], model_rows, widths=[1.2, 4.0, 1.1], font_size=9.3)

    # 添加模型参数量和版本对比说明。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 模型对比正文。
        "当前 cnn_v2 模型总参数量为 1,243,572，均为可训练参数。项目同时保留 EatingSoundCNN（cnn_v1）作为兼容基线，便于与较浅模型进行对比。cnn_v2 相比 cnn_v1 增加了卷积层深度和分类器容量，因此能够学习更细粒度的声音纹理特征。",
    )


# 定义实验设计章节生成函数。
def add_experiment_section(doc):
    # 添加实验设计一级标题。
    add_heading(doc, "3 实验设计", 1)

    # 添加实验环境二级标题。
    add_heading(doc, "3.1 实验环境", 2)

    # 添加实验环境说明正文。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 实验环境正文。
        "项目基于 Python、PyTorch 与 torchaudio 实现。训练留痕显示实验环境使用 NVIDIA GeForce RTX 3060 Laptop GPU，显存约 6.4 GB，并启用 CUDA 与 AMP 混合精度训练。混合精度训练可以降低显存占用并提升训练速度，适合本项目中批量大小为 128 的音频频谱图分类任务。",
    )

    # 定义实验环境表数据。
    env_rows = [
        # 框架行。
        ["深度学习框架", "PyTorch + torchaudio"],

        # 设备行。
        ["设备", "CUDA GPU（NVIDIA GeForce RTX 3060 Laptop GPU）"],

        # 混合精度行。
        ["混合精度", "torch.amp autocast + GradScaler"],

        # 随机种子行。
        ["随机种子", "42"],

        # 数据读取行。
        ["数据读取", "预计算 Mel 缓存 + DataLoader"],
    ]

    # 添加表 4 标题。
    add_caption(doc, "表4 实验环境")

    # 添加实验环境表。
    add_table(doc, ["项目", "说明"], env_rows, widths=[1.8, 4.5], font_size=10)

    # 添加训练策略二级标题。
    add_heading(doc, "3.2 训练策略与超参数", 2)

    # 添加训练策略说明正文。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 训练策略正文。
        "训练入口位于 train.py。训练过程使用 AdamW 优化器，初始学习率为 8e-4，权重衰减为 1e-4，并采用 CosineAnnealingLR 余弦退火学习率调度。损失函数为带类别权重的 CrossEntropyLoss，同时设置 label_smoothing=0.05，使模型在类别不均衡和相似声音类别下具有更稳定的泛化能力。",
    )

    # 定义训练超参数表数据。
    hyper_rows = [
        # 模型行。
        ["模型", "cnn_v2"],

        # 训练轮数行。
        ["训练轮数", TRAINING_RUN.get("epochs", 60)],

        # 批量大小行。
        ["批量大小", "128"],

        # 优化器行。
        ["优化器", "AdamW"],

        # 初始学习率行。
        ["初始学习率", "8e-4"],

        # 权重衰减行。
        ["权重衰减", "1e-4"],

        # 学习率调度行。
        ["学习率调度", "CosineAnnealingLR，eta_min=1e-6"],

        # 损失函数行。
        ["损失函数", "CrossEntropyLoss + 类别权重 + label_smoothing=0.05"],

        # 早停策略行。
        ["早停策略", "验证准确率连续 15 轮不提升时停止"],

        # 训练增强行。
        ["训练增强", "TimeMasking、FrequencyMasking、时间轴随机平移"],
    ]

    # 添加表 5 标题。
    add_caption(doc, "表5 训练超参数")

    # 添加训练超参数表。
    add_table(doc, ["参数", "设置"], hyper_rows, widths=[1.8, 4.5], font_size=10)

    # 添加对比实验二级标题。
    add_heading(doc, "3.3 对比实验设计", 2)

    # 添加对比实验说明正文。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 对比实验正文。
        f"项目迭代过程中保留了两份 cnn_v1 训练留痕，并在最终版本中切换到更强的 {CHECKPOINT['arch']}。对比目标是观察模型容量、双卷积模块和更充分训练轮数对验证准确率的影响。早期 cnn_v1 的参数量约为 423,796，最终 cnn_v2 参数量为 1,243,572；二者均使用 Mel 频谱图作为输入。本次刷新后的 cnn_v2 完成 {TRAINING_RUN.get('epochs', 60)} 轮训练，并在第 {CHECKPOINT['epoch']} 轮取得 {fmt_pct(EVALUATION['accuracy'])} 的最佳验证准确率。",
    )

    # 定义模型迭代对比表数据。
    comparison_rows = [
        # 第一次 cnn_v1 留痕实验。
        ["cnn_v1 留痕实验 1", "423,796", "50", "82.50%", "留痕.txt"],

        # 第二次 cnn_v1 留痕实验。
        ["cnn_v1 留痕实验 2", "423,796", "50", "76.44%", "留痕2.txt"],

        # 最终 cnn_v2 模型结果。
        ["cnn_v2 最终模型", "1,243,572", TRAINING_RUN.get("epochs", 60), fmt_pct(EVALUATION["accuracy"]), "best_model.pth"],
    ]

    # 添加表 6 标题。
    add_caption(doc, "表6 模型迭代对比")

    # 添加模型迭代表。
    add_table(doc, ["实验", "参数量", "训练轮数", "最佳验证准确率", "依据"], comparison_rows, widths=[1.7, 1.2, 1.0, 1.4, 1.2], font_size=9.5)


# 定义实验结果章节生成函数。
def add_result_section(doc):
    # 添加实验结果一级标题。
    add_heading(doc, "4 实验结果分析", 1)

    # 添加训练过程二级标题。
    add_heading(doc, "4.1 训练过程分析", 2)

    # 添加训练曲线说明正文。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 训练过程分析正文。
        f"最新训练历史记录覆盖 {TRAINING_RUN.get('epochs', 60)} 轮，训练损失和验证损失整体下降，训练准确率与验证准确率稳步上升并在后期趋于收敛。图3 中星标表示最终保存的 {CHECKPOINT['arch']} 模型最优验证点，该模型在第 {CHECKPOINT['epoch']} 轮达到 {fmt_pct(EVALUATION['accuracy'])} 的验证准确率；第 {CHECKPOINT['epoch']} 轮之后验证准确率仅在 94% 左右小幅波动，说明模型已经接近当前配置下的收敛区间。",
    )

    # 添加训练曲线图。
    add_picture(doc, "training_curves.png", 6.25, "图3 训练与验证曲线记录")

    # 添加最优模型结果二级标题。
    add_heading(doc, "4.2 最优模型结果", 2)

    # 从指标字典中取出 evaluation 部分。
    evaluation = METRICS["evaluation"]

    # 从 evaluation 中取出 classification_report。
    report = evaluation["classification_report"]

    # 定义最终结果表数据。
    result_rows = [
        # 最佳模型结构。
        ["最佳模型结构", METRICS["checkpoint"]["arch"]],

        # 最佳保存轮次。
        ["最佳保存轮次", f"第 {METRICS['checkpoint']['epoch']} 轮"],

        # 验证集样本数。
        ["验证集样本数", evaluation["val_samples"]],

        # 训练集样本数。
        ["训练集样本数", evaluation["train_samples"]],

        # 验证准确率。
        ["验证准确率", f"{evaluation['accuracy'] * 100:.2f}%"],

        # 宏平均精确率。
        ["Macro Precision", f"{report['macro avg']['precision'] * 100:.2f}%"],

        # 宏平均召回率。
        ["Macro Recall", f"{report['macro avg']['recall'] * 100:.2f}%"],

        # 宏平均 F1。
        ["Macro F1-score", f"{report['macro avg']['f1-score'] * 100:.2f}%"],

        # 加权 F1。
        ["Weighted F1-score", f"{report['weighted avg']['f1-score'] * 100:.2f}%"],
    ]

    # 添加表 7 标题。
    add_caption(doc, "表7 最优模型验证结果")

    # 添加最优模型结果表。
    add_table(doc, ["指标", "数值"], result_rows, widths=[2.2, 3.8], font_size=10)

    # 添加结果总结正文。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 结果总结正文。
        f"最终模型在验证集上的准确率为 {fmt_pct(EVALUATION['accuracy'])}，宏平均 F1 为 {fmt_pct(REPORT['macro avg']['f1-score'])}，加权平均 F1 为 {fmt_pct(REPORT['weighted avg']['f1-score'])}。这些结果说明模型在多数类别上已经具备稳定识别能力，同时类别权重与 label smoothing 对类别不均衡问题起到了缓解作用；但 jelly、burger、chocolate 等相对易混或样本较少的类别仍拉低了整体上限。",
    )

    # 添加混淆矩阵二级标题。
    add_heading(doc, "4.3 混淆矩阵分析", 2)

    # 添加混淆矩阵图片。
    add_picture(doc, "confusion_matrix.png", 6.15, "图4 最优模型验证集混淆矩阵")

    # 取误分最多的前 6 组类别。
    top_confusions = evaluation["top_confusions"][:6]

    # 把误分组合拼成一段中文说明。
    confusion_text = "、".join(
        # 每个 item 生成 “true→pred（count 条）” 格式。
        [f"{item['true']}→{item['pred']}（{item['count']} 条）" for item in top_confusions]
    )

    # 添加混淆矩阵文字分析。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 混淆矩阵分析正文。
        f"混淆矩阵显示，大多数类别的样本集中分布在主对角线附近，说明模型整体分类稳定。误分较多的组合主要包括 {confusion_text}。最新结果中最明显的混淆不再是旧报告中的 ribs→wings，而是 burger→wings（{confusion_count('burger', 'wings')} 条）、ice-cream→gummies（{confusion_count('ice-cream', 'gummies')} 条）、pickles→ice-cream（{confusion_count('pickles', 'ice-cream')} 条）以及 jelly→gummies / jelly→noodles 两组软质食物相关误分。",
    )

    # 添加单类表现和改进方向分析。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 改进方向正文。
        f"重点看 jelly、gummies、noodles：jelly 的召回率为 {class_recall('jelly')}，在 89 条验证样本中有 {confusion_count('jelly', 'gummies')} 条被判为 gummies、{confusion_count('jelly', 'noodles')} 条被判为 noodles，是当前最需要补强的类别；gummies 的召回率为 {class_recall('gummies')}、精确率为 {class_precision('gummies')}，同时吸收了来自 ice-cream 和 jelly 的样本，说明软糖类声音边界仍偏宽；noodles 的召回率达到 {class_recall('noodles')}，但精确率为 {class_precision('noodles')}，主要问题不是自身漏判，而是少量 jelly 样本被吸收到 noodles。后续可针对软质/半流质食物增加样本、加入更贴近咀嚼黏连声的增强方式，并在推理阶段使用多窗口投票降低短片段偶然误判。",
    )

    # 添加项目应用与不足二级标题。
    add_heading(doc, "4.4 项目应用与不足", 2)

    # 添加预测脚本应用说明。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 应用说明正文。
        "项目提供 predict.py 作为推理脚本，输入任意音频文件后，会复用训练阶段相同的重采样、截断/补零、Mel 频谱图提取和归一化流程，加载 best_model.pth 输出 Top-5 预测类别及置信度。该脚本保证训练与推理特征处理一致，便于后续扩展为桌面端或 Web 端音频识别应用。",
    )

    # 添加项目不足和后续改进说明。
    add_body(
        # 当前 Word 文档对象。
        doc,

        # 不足与改进正文。
        "当前项目已经保存逐轮 training_history.csv、主要评估图和 metrics.json，能够复现实验曲线与本次报告指标。主要不足在于仍缺少独立测试集和跨设备、跨噪声场景验证；此外，数据集中部分类别样本数量差异较大，jelly、gummies、noodles 等声音边界接近的类别仍存在误分。后续改进方向包括扩充相似类别样本、加入真实环境噪声增强、尝试预训练音频模型、记录更多推理案例，并对软质食物类别进行更细的错误样本回听分析。",
    )


# 定义参考资料章节生成函数。
def add_references(doc):
    # 添加参考资料一级标题。
    add_heading(doc, "参考资料与项目文件", 1)

    # 添加项目符号列表，列出数据集和主要代码文件。
    add_bullets(
        # 当前 Word 文档对象。
        doc,

        # 列表项内容。
        [
            # 数据集来源。
            "数据集：Kaggle Eating Sound Collection，https://www.kaggle.com/datasets/mashijie/eating-sound-collection?resource=download",

            # 配置文件说明。
            "项目配置文件：config.py，包含类别、音频参数、训练超参数和路径配置。",

            # 数据处理文件说明。
            "数据处理文件：dataset.py，包含 Mel 频谱图缓存、分层划分、训练增强和验证集归一化。",

            # 模型文件说明。
            "模型文件：model.py，包含 cnn_v1、cnn_v2 及 build_model 构建函数。",

            # 训练文件说明。
            "训练文件：train.py，包含训练循环、验证、AMP、学习率调度、早停和模型保存逻辑。",

            # 推理文件说明。
            "推理文件：predict.py，包含单音频文件预处理、模型加载和 Top-5 预测输出。",
        ],
    )


# 定义主函数；把所有报告章节按顺序生成。
def main():
    # 使用模板文件创建 Document 对象。
    doc = Document(str(TEMPLATE))

    # 清空模板正文，只保留页面节属性。
    clear_body(doc)

    # 设置页边距和默认字体。
    setup_styles(doc)

    # 添加封面。
    add_cover(doc)

    # 添加摘要。
    add_abstract(doc)

    # 添加数据集章节。
    add_dataset_section(doc)

    # 添加模型章节。
    add_model_section(doc)

    # 添加实验设计章节。
    add_experiment_section(doc)

    # 添加实验结果章节。
    add_result_section(doc)

    # 添加参考资料章节。
    add_references(doc)

    # 保存最终 Word 报告。
    doc.save(str(OUTPUT))

    # 打印输出文件路径，方便用户知道报告生成在哪里。
    print(OUTPUT)


# 只有直接运行 python build_report_docx.py 时才会执行主函数。
if __name__ == "__main__":
    # 调用主函数，开始生成报告。
    main()
